from __future__ import annotations

import re
from dataclasses import dataclass
from typing import List

import fitz

_ws = re.compile(r"\s+")


def normalize_text(s: str) -> str:
    s = s.lower()
    s = _ws.sub(" ", s).strip()
    return s


@dataclass
class TextBlock:
    page: int
    bbox: list
    text: str
    text_norm: str


def extract_pdf_blocks(pdf_path: str) -> List[TextBlock]:
    doc = fitz.open(pdf_path)
    blocks: List[TextBlock] = []
    for i, page in enumerate(doc):
        raw = page.get_text("blocks")
        for b in raw:
            x0, y0, x1, y1, txt = b[0], b[1], b[2], b[3], b[4]
            if not txt or not txt.strip():
                continue
            tnorm = normalize_text(txt)
            blocks.append(
                TextBlock(page=i + 1, bbox=[x0, y0, x1, y1], text=txt.strip(), text_norm=tnorm)
            )
    doc.close()
    return blocks

def extract_docx_blocks(docx_path: str) -> List[TextBlock]:
    import docx
    doc = docx.Document(docx_path)
    blocks: List[TextBlock] = []
    
    for i, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        if not txt:
            continue
        tnorm = normalize_text(txt)
        blocks.append(
            TextBlock(page=1, bbox=[0, 0, 0, 0], text=txt, text_norm=tnorm)
        )
        
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                txt = cell.text.strip()
                if not row_cells or row_cells[-1] != txt:
                    row_cells.append(txt)
            row_txt = " | ".join([t for t in row_cells if t])
            if row_txt:
                tnorm = normalize_text(row_txt)
                blocks.append(
                    TextBlock(page=1, bbox=[0, 0, 0, 0], text=row_txt, text_norm=tnorm)
                )
    return blocks


def extract_text_blocks(file_path: str) -> List[TextBlock]:
    if file_path.lower().endswith(".docx"):
        return extract_docx_blocks(file_path)
    else:
        return extract_pdf_blocks(file_path)